def format_schema(schema_info):
    """Formats the schema information into a string."""
    schema_string = ""
    for table_name, columns in schema_info.items():
        schema_string += f"Table: {table_name}\n"
        for column in columns:
            schema_string += f"  Column: {column['name']} ({column['type']}) - {column['description']}\n"
    return schema_string

def parse_schema(docx_path):
    """
    Parses a Word document containing database schema information and returns a structured representation.

    Args:
        docx_path (str): The path to the Word document.

    Returns:
        dict: A dictionary containing the schema information, with table names as keys and a list of column dictionaries as values.
             Each column dictionary contains 'name', 'type', and 'description' keys.
    """
    import docx
    try:
        doc = docx.Document(docx_path)
        table_name = None
        columns = []

        for table in doc.tables:
            # Assuming the first row of the table contains column headers
            header_cells = table.rows[0].cells
            headers = [cell.text.strip() for cell in header_cells]

            # Determine the indices of the relevant columns
            name_index = headers.index('FieldName') if 'FieldName' in headers else None
            type_index = headers.index('Type') if 'Type' in headers else None
            description_index = headers.index('Description') if 'Description' in headers else None

            # Extract data from the remaining rows
            for row in table.rows[1:]:
                cells = row.cells
                column = {}

                if name_index is not None and len(cells) > name_index:
                    column['name'] = cells[name_index].text.strip()
                if type_index is not None and len(cells) > type_index:
                    column['type'] = cells[type_index].text.strip()
                if description_index is not None and len(cells) > description_index:
                    column['description'] = cells[description_index].text.strip()

                if column:
                    columns.append(column)

        # Extract table name from the document (assuming it's in the first paragraph)
        if doc.paragraphs:
            table_name = doc.paragraphs[0].text.split('Table Name')[-1].strip()

        if table_name:
            return {table_name: columns}
        else:
            return None

    except Exception as e:
        print(f"Error parsing schema from {docx_path}: {e}")
        return None

if __name__ == '__main__':
    # Example usage
    schema = parse_schema('src/DB_Schema_Test/員工薪資.docx')
    if schema:
        print(schema)
